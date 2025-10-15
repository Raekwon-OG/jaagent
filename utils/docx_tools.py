"""
DOCX document processing and PDF conversion utilities
"""
import logging
import re
import tempfile
from typing import Dict, List, Optional, Tuple
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from docx2pdf import convert
    PDF_CONVERSION_AVAILABLE = True
except ImportError:
    PDF_CONVERSION_AVAILABLE = False

from config import settings

logger = logging.getLogger(__name__)

class DocxProcessor:
    """Processes DOCX documents for resume and cover letter generation"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        self.pdf_conversion_enabled = PDF_CONVERSION_AVAILABLE and settings.GENERATE_PDF
    
    def extract_text_from_docx(self, docx_path: Path) -> str:
        """Extract text content from DOCX file"""
        
        try:
            doc = Document(str(docx_path))
            
            # Extract text from all paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            full_text = '\n'.join(text_content)
            logger.debug(f"Extracted {len(full_text)} characters from {docx_path}")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Error extracting text from {docx_path}: {e}")
            raise
    
    def extract_resume_sections(self, docx_path: Path) -> Dict[str, str]:
        """Extract specific sections from resume DOCX"""
        
        try:
            doc = Document(str(docx_path))
            
            sections = {
                'summary': '',
                'experience': '',
                'skills': '',
                'education': '',
                'other': ''
            }
            
            current_section = 'other'
            section_content = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                
                if not text:
                    continue
                
                # Detect section headers
                text_lower = text.lower()
                
                # Save previous section content
                if section_content and current_section:
                    sections[current_section] = '\n'.join(section_content)
                    section_content = []
                
                # Identify new section
                if any(keyword in text_lower for keyword in ['summary', 'profile', 'objective']):
                    current_section = 'summary'
                elif any(keyword in text_lower for keyword in ['experience', 'employment', 'work history', 'career']):
                    current_section = 'experience'
                elif any(keyword in text_lower for keyword in ['skills', 'technical', 'competencies']):
                    current_section = 'skills'
                elif any(keyword in text_lower for keyword in ['education', 'academic', 'qualifications']):
                    current_section = 'education'
                else:
                    # Add to current section
                    section_content.append(text)
            
            # Save final section
            if section_content and current_section:
                sections[current_section] = '\n'.join(section_content)
            
            logger.info(f"Extracted resume sections from {docx_path}")
            return sections
            
        except Exception as e:
            logger.error(f"Error extracting resume sections from {docx_path}: {e}")
            raise
    
    def create_tailored_resume(self, base_docx_path: Path, tailored_summary: str, 
                              tailored_experience: str, output_path: Path) -> Path:
        """Create tailored resume by updating summary and experience sections"""
        
        try:
            # Load base document
            doc = Document(str(base_docx_path))
            
            # Track which sections we've updated
            summary_updated = False
            experience_updated = False
            
            # Process paragraphs to find and update sections
            current_section = None
            paragraphs_to_remove = []
            
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip().lower()
                
                if not text:
                    continue
                
                # Detect section headers
                if any(keyword in text for keyword in ['summary', 'profile', 'objective']):
                    current_section = 'summary'
                    continue
                elif any(keyword in text for keyword in ['experience', 'employment', 'work history', 'career']):
                    current_section = 'experience'
                    continue
                elif any(keyword in text for keyword in ['skills', 'technical', 'education']):
                    current_section = None
                    continue
                
                # Mark paragraphs for replacement in target sections
                if current_section == 'summary' and not summary_updated:
                    # Replace summary content
                    paragraph.clear()
                    paragraph.add_run(tailored_summary)
                    summary_updated = True
                    
                    # Mark subsequent summary paragraphs for removal
                    for j in range(i + 1, len(doc.paragraphs)):
                        next_para = doc.paragraphs[j]
                        next_text = next_para.text.strip().lower()
                        if any(keyword in next_text for keyword in ['experience', 'employment', 'skills', 'education']):
                            break
                        if next_para.text.strip():
                            paragraphs_to_remove.append(j)
                
                elif current_section == 'experience' and not experience_updated:
                    # Replace experience content
                    paragraph.clear()
                    paragraph.add_run(tailored_experience)
                    experience_updated = True
                    
                    # Mark subsequent experience paragraphs for removal
                    for j in range(i + 1, len(doc.paragraphs)):
                        next_para = doc.paragraphs[j]
                        next_text = next_para.text.strip().lower()
                        if any(keyword in next_text for keyword in ['skills', 'education', 'certifications']):
                            break
                        if next_para.text.strip():
                            paragraphs_to_remove.append(j)
            
            # Remove marked paragraphs (in reverse order to maintain indices)
            for i in sorted(paragraphs_to_remove, reverse=True):
                if i < len(doc.paragraphs):
                    p = doc.paragraphs[i]._element
                    p.getparent().remove(p)
            
            # Save tailored document
            doc.save(str(output_path))
            logger.info(f"Created tailored resume: {output_path}")
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error creating tailored resume: {e}")
            raise
    
    def create_cover_letter(self, cover_letter_text: str, output_path: Path) -> Path:
        """Create cover letter DOCX from text"""
        
        try:
            # Create new document
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Split cover letter into lines and add as paragraphs
            lines = cover_letter_text.split('\n')
            
            for line in lines:
                line = line.strip()
                if line:
                    paragraph = doc.add_paragraph(line)
                    
                    # Format dates and addresses
                    if self._is_date_line(line):
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif self._is_address_line(line):
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    # Add empty paragraph for spacing
                    doc.add_paragraph()
            
            # Save document
            doc.save(str(output_path))
            logger.info(f"Created cover letter: {output_path}")
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error creating cover letter: {e}")
            raise
    
    def _is_date_line(self, line: str) -> bool:
        """Check if line contains a date"""
        date_patterns = [
            r'\d{1,2}/\d{1,2}/\d{4}',
            r'\d{1,2}-\d{1,2}-\d{4}',
            r'[A-Za-z]+ \d{1,2}, \d{4}',
            r'\d{1,2} [A-Za-z]+ \d{4}'
        ]
        
        for pattern in date_patterns:
            if re.search(pattern, line):
                return True
        return False
    
    def _is_address_line(self, line: str) -> bool:
        """Check if line is part of an address"""
        address_indicators = [
            'street', 'avenue', 'road', 'drive', 'lane', 'blvd', 'ave', 'st',
            'suite', 'apt', 'unit', 'floor', 'building'
        ]
        
        line_lower = line.lower()
        return any(indicator in line_lower for indicator in address_indicators)
    
    def convert_to_pdf(self, docx_path: Path, pdf_path: Optional[Path] = None) -> Optional[Path]:
        """Convert DOCX to PDF"""
        
        if not self.pdf_conversion_enabled:
            logger.warning("PDF conversion is disabled or docx2pdf not available")
            return None
        
        if pdf_path is None:
            pdf_path = docx_path.with_suffix('.pdf')
        
        try:
            convert(str(docx_path), str(pdf_path))
            logger.info(f"Converted to PDF: {pdf_path}")
            return pdf_path
            
        except Exception as e:
            logger.error(f"Error converting {docx_path} to PDF: {e}")
            
            # Try alternative conversion method if available
            return self._alternative_pdf_conversion(docx_path, pdf_path)
    
    def _alternative_pdf_conversion(self, docx_path: Path, pdf_path: Path) -> Optional[Path]:
        """Alternative PDF conversion method using system tools"""
        
        try:
            import subprocess
            import sys
            
            # Try LibreOffice conversion (if available)
            if sys.platform.startswith('linux') or sys.platform == 'darwin':
                cmd = ['libreoffice', '--headless', '--convert-to', 'pdf', 
                       '--outdir', str(pdf_path.parent), str(docx_path)]
                
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                
                if result.returncode == 0 and pdf_path.exists():
                    logger.info(f"PDF converted using LibreOffice: {pdf_path}")
                    return pdf_path
            
        except Exception as e:
            logger.debug(f"Alternative PDF conversion failed: {e}")
        
        return None
    
    def create_job_documents(self, job_folder: Path, tailored_resume_text: str, 
                           tailored_cover_letter_text: str, base_resume_path: Path) -> Dict[str, Path]:
        """Create all job application documents"""
        
        document_paths = {}
        
        try:
            # Create tailored resume DOCX
            resume_docx_path = job_folder / f"{settings.RESUME_FILENAME}.docx"
            
            # For simplicity, create new resume from text rather than modifying existing
            # In production, you might want to preserve formatting from base resume
            self._create_resume_from_text(tailored_resume_text, resume_docx_path)
            document_paths['resume_docx'] = resume_docx_path
            
            # Create cover letter DOCX
            cover_letter_docx_path = job_folder / f"{settings.COVER_LETTER_FILENAME}.docx"
            self.create_cover_letter(tailored_cover_letter_text, cover_letter_docx_path)
            document_paths['cover_letter_docx'] = cover_letter_docx_path
            
            # Convert to PDF if enabled
            if self.pdf_conversion_enabled:
                resume_pdf_path = self.convert_to_pdf(resume_docx_path)
                if resume_pdf_path:
                    document_paths['resume_pdf'] = resume_pdf_path
                
                cover_letter_pdf_path = self.convert_to_pdf(cover_letter_docx_path)
                if cover_letter_pdf_path:
                    document_paths['cover_letter_pdf'] = cover_letter_pdf_path
            
            logger.info(f"Created {len(document_paths)} job documents")
            return document_paths
            
        except Exception as e:
            logger.error(f"Error creating job documents: {e}")
            raise
    
    def _create_resume_from_text(self, resume_text: str, output_path: Path):
        """Create a resume DOCX from plain text"""
        
        try:
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.8)
                section.bottom_margin = Inches(0.8)
                section.left_margin = Inches(0.8)
                section.right_margin = Inches(0.8)
            
            # Split text into paragraphs
            paragraphs = resume_text.split('\n')
            
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    paragraph = doc.add_paragraph(para_text)
                    
                    # Format section headers (all caps or title case with colons)
                    if (para_text.isupper() and len(para_text) < 50) or para_text.endswith(':'):
                        paragraph.runs[0].bold = True
                        paragraph.runs[0].font.size = 12
                else:
                    # Add spacing between sections
                    doc.add_paragraph()
            
            doc.save(str(output_path))
            
        except Exception as e:
            logger.error(f"Error creating resume from text: {e}")
            raise
    
    def validate_docx_setup(self) -> List[str]:
        """Validate DOCX processing setup"""
        
        issues = []
        
        if not DOCX_AVAILABLE:
            issues.append("python-docx not installed. Run: pip install python-docx")
        
        if settings.GENERATE_PDF and not PDF_CONVERSION_AVAILABLE:
            issues.append("docx2pdf not installed. PDF generation disabled. Run: pip install docx2pdf")
        
        # Test document creation
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=True) as tmp_file:
                doc = Document()
                doc.add_paragraph("Test document")
                doc.save(tmp_file.name)
        except Exception as e:
            issues.append(f"Cannot create DOCX documents: {e}")
        
        # Test PDF conversion if enabled
        if self.pdf_conversion_enabled:
            try:
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx, \
                     tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_pdf:
                    
                    doc = Document()
                    doc.add_paragraph("Test PDF conversion")
                    doc.save(tmp_docx.name)
                    
                    # Try conversion
                    result_pdf = self.convert_to_pdf(Path(tmp_docx.name), Path(tmp_pdf.name))
                    
                    if not result_pdf or not result_pdf.exists():
                        issues.append("PDF conversion not working")
                    
                    # Cleanup
                    Path(tmp_docx.name).unlink(missing_ok=True)
                    Path(tmp_pdf.name).unlink(missing_ok=True)
                    
            except Exception as e:
                issues.append(f"PDF conversion test failed: {e}")
        
        return issues


# Factory function for easy access
def create_docx_processor() -> DocxProcessor:
    """Create a configured DOCX processor instance"""
    return DocxProcessor()


# Global processor instance
_global_docx_processor: Optional[DocxProcessor] = None

def get_global_docx_processor() -> DocxProcessor:
    """Get or create the global DOCX processor instance"""
    global _global_docx_processor
    if _global_docx_processor is None:
        _global_docx_processor = create_docx_processor()
    return _global_docx_processor


# Convenience functions
def extract_resume_text(docx_path: Path) -> str:
    """Quick function to extract text from resume"""
    processor = get_global_docx_processor()
    return processor.extract_text_from_docx(docx_path)

def create_documents_from_text(job_folder: Path, resume_text: str, cover_letter_text: str) -> Dict[str, Path]:
    """Quick function to create documents from text"""
    processor = get_global_docx_processor()
    
    # Create documents manually since we don't have base resume
    docs = {}
    
    # Resume
    resume_docx = job_folder / f"{settings.RESUME_FILENAME}.docx"
    processor._create_resume_from_text(resume_text, resume_docx)
    docs['resume_docx'] = resume_docx
    
    # Cover letter
    cover_letter_docx = job_folder / f"{settings.COVER_LETTER_FILENAME}.docx"
    processor.create_cover_letter(cover_letter_text, cover_letter_docx)
    docs['cover_letter_docx'] = cover_letter_docx
    
    # PDFs if enabled
    if processor.pdf_conversion_enabled:
        resume_pdf = processor.convert_to_pdf(resume_docx)
        if resume_pdf:
            docs['resume_pdf'] = resume_pdf
            
        cover_letter_pdf = processor.convert_to_pdf(cover_letter_docx)
        if cover_letter_pdf:
            docs['cover_letter_pdf'] = cover_letter_pdf
    
    return docs